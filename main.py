from fastapi import FastAPI, HTTPException, Request, Form
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
import os
import datetime
import subprocess
import re

# Inicializa a aplicação FastAPI
app = FastAPI()

# Configuração dos templates e arquivos estáticos
templates = Jinja2Templates(directory="templates")
app.mount("/static", StaticFiles(directory="static"), name="static")

# Diretório onde os resumos serão salvos
OUTPUT_DIR = "resumos"
os.makedirs(OUTPUT_DIR, exist_ok=True)  # Cria o diretório se não existir

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """
    Rota da página inicial.
    Retorna o template "index.html".
    """
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/summarize/")
async def summarize_text(request: Request, text: str = Form(...)):
    """
    Endpoint para resumir um texto e salvar o resultado em um arquivo Word.
    """
    try:
        # Executa o modelo para gerar um título e um resumo
        title, summary = run_llama(text)

        # Salva o resumo em um arquivo Word e obtém o caminho do arquivo
        file_path = save_to_word(title, summary)

        return templates.TemplateResponse(
            "index.html",
            {"request": request, "summary": summary, "file_path": file_path},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))  # Retorna erro HTTP 500 em caso de falha

@app.get("/download/{file_name}")
async def download_file(file_name: str):
    """
    Permite o download de um arquivo Word gerado pelo sistema.
    """
    file_path = os.path.join(OUTPUT_DIR, file_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado.")
    return FileResponse(
        file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def run_llama(text):
    """
    Executa o modelo LLaMA via Ollama para gerar um título e um resumo do texto fornecido.
    """
    try:
        # Comando para chamar o modelo LLaMA via terminal
        command = [
            "ollama", "run", "llama3.2:1b", 
            f"Resuma o seguinte texto e gere um título curto. Use este formato: 'Título: <título gerado>\nResumo: <resumo gerado>'\n\n{text}"
        ]
        result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding="utf-8")

        # Verifica se a execução foi bem-sucedida
        if result.returncode != 0:
            raise Exception(f"Ollama error: {result.stderr.strip()}")

        # Captura a saída do comando
        output = result.stdout.strip()
        print(f"Saída do comando: {output}")

        # Extrai título e resumo do formato esperado
        match = re.search(r"Título:\s*(.*?)\s*\nResumo:\s*(.*)", output, re.DOTALL | re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            summary = match.group(2).strip()
        else:
            # Se o formato esperado não for encontrado, tenta dividir manualmente
            lines = output.split("\n", 1)
            if len(lines) > 1 and len(lines[0]) < 50:  # Assume que títulos são curtos
                title = lines[0].strip()
                summary = lines[1].strip()
            else:
                raise Exception("Formato de saída inesperado do modelo.")

        return title, summary
    except Exception as e:
        raise Exception(f"Erro ao executar o modelo: {str(e)}")

def sanitize_filename(title):
    """
    Remove caracteres inválidos do título para que ele possa ser usado como nome de arquivo.
    """
    return re.sub(r'[<>:"/\\|?*]', '', title)  # Remove caracteres não permitidos no Windows

def save_to_word(title, summary):
    """
    Salva o título e o resumo em um arquivo Word.
    """
    # Certifica-se de que o diretório de saída existe
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    # Cria um novo documento Word
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(summary)

    # Sanitiza o título para criar um nome de arquivo válido
    sanitized_title = sanitize_filename(title)

    # Define o nome do arquivo com base no título e data/hora
    file_name = f"{sanitized_title}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)

    # Salva o documento no diretório especificado
    doc.save(file_path)
    return file_name
